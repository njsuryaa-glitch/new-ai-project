import io
from uuid import UUID
from fastapi import HTTPException, status
from pypdf import PdfReader
import docx

from app.db.models import Document, DocumentChunk
from app.repositories.document_repository import DocumentRepository
from app.services.embedding_service import EmbeddingService
from app.core.config import settings
from app.core.logging import get_logger

logger = get_logger(__name__)


class DocumentService:
    def __init__(
        self,
        document_repository: DocumentRepository,
        embedding_service: EmbeddingService
    ) -> None:
        self.repo = document_repository
        self.embedding_service = embedding_service

    def extract_text(self, filename: str, content: bytes) -> str:
        """
        Extracts plain text from the uploaded file based on its extension.
        Supports PDF, DOCX, and TXT files.
        """
        ext = filename.split(".")[-1].lower()
        text = ""

        if ext == "pdf":
            try:
                reader = PdfReader(io.BytesIO(content))
                pages_text = []
                for idx, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
                text = "\n".join(pages_text)
            except Exception as e:
                logger.error(f"Error parsing PDF file {filename}: {e}")
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Failed to parse PDF file: {str(e)}"
                )

        elif ext == "docx":
            try:
                doc = docx.Document(io.BytesIO(content))
                text = "\n".join([para.text for para in doc.paragraphs])
            except Exception as e:
                logger.error(f"Error parsing DOCX file {filename}: {e}")
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Failed to parse DOCX file: {str(e)}"
                )

        elif ext == "txt":
            try:
                text = content.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    text = content.decode("latin-1")
                except Exception as e:
                    logger.error(f"Error decoding TXT file {filename}: {e}")
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail="Failed to decode TXT file. Ensure it is encoded in UTF-8 or Latin-1."
                    )
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file type: {ext}"
            )

        if not text.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Extracted text is empty. The document may be empty or contain only image data."
            )

        return text

    def chunk_text(self, text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
        """
        Chunks text into sections using a character-based approach.
        Attempts to split at paragraphs, newlines, or spaces to preserve word integrity.
        """
        if chunk_overlap >= chunk_size:
            chunk_overlap = chunk_size // 2

        chunks = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = min(start + chunk_size, text_len)

            if end < text_len:
                # Search backward for natural separators to avoid cutting words
                break_idx = -1
                for sep in ("\n\n", "\n", " "):
                    idx = text.rfind(sep, start + chunk_overlap, end)
                    if idx != -1:
                        # Split after the separator
                        break_idx = idx + len(sep)
                        break
                if break_idx != -1:
                    end = break_idx

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            # Shift window forward
            start = end - chunk_overlap
            
            # Avoid infinite loops or exceeding bounds
            if start < 0 or end == text_len:
                break

        return chunks

    async def process_document(self, filename: str, content: bytes) -> tuple[Document, int]:
        """
        Extracts, chunks, embeds, and saves the document contents.
        This runs as a transaction using the session's flush/commit cycle.
        """
        # 1. Parse content
        extracted_text = self.extract_text(filename, content)

        # 2. Chunk text
        chunks = self.chunk_text(
            extracted_text,
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP
        )

        if not chunks:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No chunks could be created from document text."
            )

        # 3. Create document record
        doc = Document(filename=filename)
        await self.repo.save_document(doc)

        # 4. Generate embeddings for all chunks in batch
        try:
            embeddings = await self.embedding_service.get_embeddings(chunks)
        except Exception as e:
            logger.error(f"Failed to generate embeddings for document {filename}: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Internal AI service failed to generate embeddings."
            )

        # 5. Create chunk records
        chunk_records = []
        for idx, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
            chunk_records.append(
                DocumentChunk(
                    document_id=doc.id,
                    chunk_text=chunk_text,
                    embedding=embedding,
                    chunk_index=idx
                )
            )

        # 6. Save chunks and commit the full transaction
        await self.repo.save_chunks(chunk_records)
        await self.repo.commit()

        logger.info(f"Processed document {filename} (ID: {doc.id}). Created {len(chunk_records)} chunks.")
        return doc, len(chunk_records)
